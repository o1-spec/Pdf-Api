import pyotp
from PIL import Image
import os
import time
from pdf2image import convert_from_path
import zipfile
import subprocess
from pdf2docx import Converter
from docx import Document
import pandas as pd
import fitz

def generate_otp():
    totp = pyotp.TOTP(pyotp.random_base32(), interval=300)
    return totp.now()

def verify_otp(otp, user_otp):
    return otp == user_otp

def convert_images_to_pdf(image_files, output_pdf_path):
    timestamp = str(int(time.time()))
    output_pdf_path = output_pdf_path.replace('output.pdf', f'output_{timestamp}.pdf')
    images = []
    for image_file in image_files:
        img = Image.open(image_file)
        if img.mode != 'RGB':
            img = img.convert('RGB')
        images.append(img)

    if images:
        images[0].save(output_pdf_path, save_all=True, append_images=images[1:])
        return output_pdf_path

def convert_pdf_to_images(pdf_file_path, output_folder):
    try:
        images = convert_from_path(pdf_file_path)
        output_files = []
        for i, image in enumerate(images):
            output_path = os.path.join(output_folder, f'page_{i + 1}.png')
            image.save(output_path, 'PNG')
            output_files.append(output_path)
            
        if len(images) == 1:
            return {"type": "single", "file_path": output_files[0]}
        else:
            zip_file_path = os.path.join(output_folder, 'converted_images.zip')
            with zipfile.ZipFile(zip_file_path, 'w') as zipf:
                for file in output_files:
                    zipf.write(file, os.path.basename(file))
            return {"type": "zip", "file_path": zip_file_path}
        
    except Exception as e:
        print("An unexpected error occurred:", e)
        return None

def convert_word_to_pdf(word_file_path, output_folder):
    output_pdf_path = os.path.join(output_folder, os.path.splitext(os.path.basename(word_file_path))[0] + ".pdf")

    try:
        command = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", output_folder, word_file_path]
        subprocess.run(command, check=True)
        return output_pdf_path
    except Exception as e:
        print("Error converting Word to PDF:", e)
        return None

def convert_pdf_to_word(pdf_file_path, output_folder):
    output_docx_path = os.path.join(output_folder, os.path.splitext(os.path.basename(pdf_file_path))[0] + ".docx")

    try:
        converter = Converter(pdf_file_path)
        converter.convert(output_docx_path, start=0, end=None)
        converter.close()
        return output_docx_path
    except Exception as e:
        print("Error converting PDF to Word:", e)
        return None

def convert_word_to_excel(word_file_path, output_folder):
    try:
        doc = Document(word_file_path)
        data = []

        # Read paragraphs from Word
        for para in doc.paragraphs:
            if para.text.strip():
                data.append([para.text.strip()])

        # Convert to DataFrame
        df = pd.DataFrame(data, columns=["Word Content"])

        # Save to Excel
        output_excel_path = os.path.join(output_folder, "converted_word.xlsx")
        df.to_excel(output_excel_path, index=False)
        return output_excel_path

    except Exception as e:
        print("Error converting Word to Excel:", e)
        return None

def convert_excel_to_word(excel_file_path, output_folder):
    try:
        # Read Excel file
        df = pd.read_excel(excel_file_path)

        # Create Word document
        doc = Document()
        doc.add_heading("Excel Data to Word", level=1)

        # Add a table to Word
        table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])

        # Add headers
        for j, column_name in enumerate(df.columns):
            table.cell(0, j).text = column_name

        # Add rows
        for i, row in df.iterrows():
            for j, value in enumerate(row):
                table.cell(i + 1, j).text = str(value)

        # Save Word file
        output_word_path = os.path.join(output_folder, "converted_excel.docx")
        doc.save(output_word_path)
        return output_word_path

    except Exception as e:
        print("Error converting Excel to Word:", e)
        return None
    
def compress_pdf(input_pdf_path, output_pdf_path, quality=60):
    doc = fitz.open(input_pdf_path)
    
    for page in doc:
        for img in page.get_images(full=True):
            xref = img[0]  # Image reference
            pix = fitz.Pixmap(doc, xref)  

            if pix.n > 4:  # Convert CMYK to RGB
                pix = fitz.Pixmap(fitz.csRGB, pix)

            new_image = pix.pil_save(quality=quality)  # Reduce quality
            doc.update_image(xref, stream=new_image)

    doc.save(output_pdf_path)
    doc.close()
    return output_pdf_path

def merge_pdfs(pdf_list, output_pdf_path):
    """Merge multiple PDFs into one."""
    merger = PyPDF2.PdfMerger()

    for pdf in pdf_list:
        merger.append(pdf)

    merger.write(output_pdf_path)
    merger.close()
    
    return output_pdf_path

def split_pdf(input_pdf_path, output_folder):
    """Split a PDF into individual pages and save each as a separate file."""
    os.makedirs(output_folder, exist_ok=True)

    with open(input_pdf_path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        num_pages = len(reader.pages)

        split_files = []
        for i in range(num_pages):
            writer = PyPDF2.PdfWriter()
            writer.add_page(reader.pages[i])

            output_pdf_path = os.path.join(output_folder, f"page_{i + 1}.pdf")
            with open(output_pdf_path, "wb") as output_file:
                writer.write(output_file)

            split_files.append(output_pdf_path)

    return split_files